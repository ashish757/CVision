import os
import time
import logging
from typing import Tuple, Dict, Any
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass


class TextExtractor:
    """Text extraction utility class for PDF and DOCX files"""

    # Thread pool for non-blocking extraction
    _executor = ThreadPoolExecutor(max_workers=4)

    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """
        Get basic file information

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with file information
        """
        path = Path(file_path)

        if not path.exists():
            raise TextExtractionError(f"File does not exist: {file_path}")

        file_stat = path.stat()
        file_extension = path.suffix.lower()

        return {
            "file_name": path.name,
            "file_extension": file_extension,
            "file_size": file_stat.st_size,
            "file_type": file_extension[1:] if file_extension else "unknown"
        }

    @staticmethod
    def _extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file using pdfplumber

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            logger.info(f"[UTIL] Starting PDF text extraction: {file_path}")

            extracted_text = []

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"[UTIL] PDF has {total_pages} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Extract text from page
                        page_text = page.extract_text()

                        if page_text:
                            # Clean up the text
                            page_text = page_text.strip()
                            if page_text:
                                extracted_text.append(page_text)
                                logger.info(f"[UTIL] Extracted {len(page_text)} characters from page {page_num}/{total_pages}")

                    except Exception as e:
                        logger.warning(f"[UTIL] Failed to extract text from page {page_num}: {str(e)}")
                        continue

            # Join all pages with double newlines
            full_text = "\n\n".join(extracted_text)

            if not full_text.strip():
                logger.error(f"[UTIL] No text content found in PDF: {file_path}")
                raise TextExtractionError("No text content found in PDF")

            logger.info(f"[UTIL] PDF extraction completed successfully. Total characters: {len(full_text)}")
            return full_text

        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            logger.error(f"[UTIL] PDF extraction failed: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def _extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            logger.info(f"[UTIL] Starting DOCX text extraction: {file_path}")

            doc = Document(file_path)
            extracted_text = []

            # Extract text from all paragraphs
            paragraph_count = len(doc.paragraphs)
            logger.info(f"[UTIL] DOCX has {paragraph_count} paragraphs")

            for i, paragraph in enumerate(doc.paragraphs, 1):
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    extracted_text.append(paragraph_text)

            # Extract text from tables
            table_count = len(doc.tables)
            logger.info(f"[UTIL] DOCX has {table_count} tables")

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        extracted_text.append(" | ".join(row_text))

            # Join all text with newlines
            full_text = "\n".join(extracted_text)

            if not full_text.strip():
                logger.error(f"[UTIL] No text content found in DOCX: {file_path}")
                raise TextExtractionError("No text content found in DOCX")

            logger.info(f"[UTIL] DOCX extraction completed successfully. Total characters: {len(full_text)}")
            return full_text

        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            logger.error(f"[UTIL] DOCX extraction failed: {str(e)}")
            raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")

    @classmethod
    def _extract_text_sync(cls, file_path: str, file_type: str) -> str:
        """
        Synchronous text extraction dispatcher

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf or docx)

        Returns:
            Extracted text content
        """
        if file_type == "pdf":
            return cls._extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return cls._extract_text_from_docx(file_path)
        else:
            raise TextExtractionError(f"Unsupported file type: {file_type}")

    @classmethod
    async def extract_text(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Asynchronously extract text from a file

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (extracted_text, file_info)

        Raises:
            TextExtractionError: If extraction fails
        """
        start_time = time.time()

        # Get file information
        file_info = cls.get_file_info(file_path)
        file_type = file_info["file_type"]

        logger.info(f"[UTIL] Starting text extraction for {file_type.upper()} file: {file_path}")
        logger.info(f"[UTIL] File size: {file_info['file_size']} bytes")

        try:
            # Run extraction in thread pool to avoid blocking event loop
            loop = asyncio.get_event_loop()
            logger.info(f"[UTIL] Running extraction in thread pool...")
            extracted_text = await loop.run_in_executor(
                cls._executor,
                cls._extract_text_sync,
                file_path,
                file_type
            )

            # Calculate processing time
            processing_time = time.time() - start_time
            file_info["processing_time_seconds"] = round(processing_time, 3)
            file_info["text_length"] = len(extracted_text)

            logger.info(f"[UTIL] Text extraction completed in {processing_time:.3f} seconds")
            logger.info(f"[UTIL] Extracted {len(extracted_text)} characters")

            return extracted_text, file_info

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"[UTIL] Text extraction failed after {processing_time:.3f} seconds: {str(e)}")
            raise

    @staticmethod
    def validate_extracted_text(text: str, min_length: int = 10) -> bool:
        """
        Validate extracted text quality

        Args:
            text: Extracted text to validate
            min_length: Minimum required length

        Returns:
            True if text is valid
        """
        if not text or not isinstance(text, str):
            return False

        # Check minimum length
        if len(text.strip()) < min_length:
            return False

        # Check if text is not just whitespace or special characters
        cleaned_text = ''.join(char for char in text if char.isalnum() or char.isspace())
        if len(cleaned_text.strip()) < min_length:
            return False

        return True

    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Clean and normalize extracted text

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            # Strip whitespace and remove empty lines
            cleaned_line = line.strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)

        # Join with single newlines and limit consecutive newlines
        cleaned_text = '\n'.join(cleaned_lines)

        # Replace multiple consecutive newlines with double newline
        import re
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)

        return cleaned_text.strip()


# Convenience functions for direct use
async def extract_text_from_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from a file (convenience function)

    Args:
        file_path: Path to the file

    Returns:
        Tuple of (extracted_text, file_info)
    """
    return await TextExtractor.extract_text(file_path)


def validate_file_for_extraction(file_path: str) -> bool:
    """
    Validate if a file can be processed for text extraction

    Args:
        file_path: Path to the file

    Returns:
        True if file is valid for extraction
    """
    try:
        file_info = TextExtractor.get_file_info(file_path)
        return file_info["file_type"] in ["pdf", "docx"]
    except Exception:
        return False
